"""Generate a professional project report DOCX for IPCMS.

Output: IPCMS_Project_Report.docx
"""

from datetime import date
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Color palette (mirrors the deck) ─────────────────────────────────────
PRIMARY   = RGBColor(0x0E, 0x4F, 0x8A)
ACCENT    = RGBColor(0x14, 0xB8, 0xA6)
SECONDARY = RGBColor(0x06, 0x6F, 0xA8)
PURPLE    = RGBColor(0x7C, 0x3A, 0xED)
PINK      = RGBColor(0xDB, 0x27, 0x77)
SLATE     = RGBColor(0x0F, 0x17, 0x2A)
BODY      = RGBColor(0x33, 0x4E, 0x68)
MUTED     = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BG  = RGBColor(0xF1, 0xF5, 0xF9)
SOFT_BG   = RGBColor(0xE0, 0xF2, 0xFE)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DIVIDER   = RGBColor(0xCB, 0xD5, 0xE1)


# ── Helpers ──────────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="CBD5E1", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_page_number(doc):
    """Add page number in footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.text = "PAGE"
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_heading(doc, text, level=1, color=PRIMARY):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(8)
    run = h.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(28)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False,
             color=BODY, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.3
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text, *, size=11, color=BODY, indent=0.25):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent + 0.1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_kv_table(doc, rows, col_widths=(2.0, 5.0), header_color=None):
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for r_idx, (k, v) in enumerate(rows):
        row = table.rows[r_idx]
        row.cells[0].width = Inches(col_widths[0])
        row.cells[1].width = Inches(col_widths[1])
        # Key
        kc = row.cells[0]
        kc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        kp = kc.paragraphs[0]
        kp.paragraph_format.space_after = Pt(2)
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(10.5)
        kr.font.color.rgb = PRIMARY
        kr.font.name = "Calibri"
        # Value
        vc = row.cells[1]
        vc.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_after = Pt(2)
        vr = vp.add_run(v)
        vr.font.size = Pt(10.5)
        vr.font.color.rgb = BODY
        vr.font.name = "Calibri"
        # Borders
        for c in (kc, vc):
            set_cell_borders(c)
        if header_color and r_idx == 0:
            shade_cell(kc, header_color)
            shade_cell(vc, header_color)
    return table


def add_grid_table(doc, header, rows, col_widths, header_fill="0E4F8A",
                   header_color=WHITE, stripe="F1F5F9"):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(header):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shade_cell(cell, header_fill)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = header_color
        r.font.name = "Calibri"
    # Data
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for i, val in enumerate(row):
            cell = row_cells[i]
            cell.width = Inches(col_widths[i])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if r_idx % 2 == 0:
                shade_cell(cell, "FFFFFF")
            else:
                shade_cell(cell, stripe)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = BODY
            r.font.name = "Calibri"
    return table


def add_callout(doc, title, body, fill="E0F2FE", border="14B8A6"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    set_cell_borders(cell, color=border, size="8")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(4)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = PRIMARY
    r1.font.name = "Calibri"
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = BODY
    r2.font.name = "Calibri"


def add_horizontal_rule(doc, color="CBD5E1"):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ── Build the document ───────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Default font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Footer with page number
add_page_number(doc)

# ── COVER ────────────────────────────────────────────────────────────────
cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
cover_title.paragraph_format.space_before = Pt(80)
r = cover_title.add_run("INTEGRATED PATIENT")
r.bold = True
r.font.size = Pt(38)
r.font.color.rgb = PRIMARY
r.font.name = "Calibri"

cover_title2 = doc.add_paragraph()
cover_title2.paragraph_format.space_after = Pt(0)
r = cover_title2.add_run("CARE MANAGEMENT SYSTEM")
r.bold = True
r.font.size = Pt(38)
r.font.color.rgb = PRIMARY
r.font.name = "Calibri"

accent = doc.add_paragraph()
accent.paragraph_format.space_after = Pt(8)
r = accent.add_run("— IP CMS —")
r.bold = True
r.font.size = Pt(18)
r.font.color.rgb = ACCENT
r.font.name = "Calibri"

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(40)
r = subtitle.add_run("A full-stack AI-assisted healthcare platform for patients, "
                      "doctors, and administrators.")
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = MUTED
r.font.name = "Calibri"

# Pill labels
pill = doc.add_paragraph()
r = pill.add_run("STREAMLIT  •  PYTHON  •  MySQL  •  LANGCHAIN  •  GROQ  •  GEMINI")
r.bold = True
r.font.size = Pt(10)
r.font.color.rgb = SLATE
r.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(40)

# Cover meta box
meta_table = doc.add_table(rows=5, cols=2)
meta = [
    ("Project Title", "Integrated Patient Care Management System (IPCMS)"),
    ("Author", "Nikhil"),
    ("Date", date.today().strftime("%B %Y")),
    ("Version", "1.0"),
    ("Domain", "Healthcare · Web Application · AI-Assisted Clinical Workflow"),
]
for i, (k, v) in enumerate(meta):
    row = meta_table.rows[i]
    row.cells[0].width = Inches(1.8)
    row.cells[1].width = Inches(4.8)
    kc = row.cells[0]
    vc = row.cells[1]
    shade_cell(kc, "F1F5F9")
    shade_cell(vc, "FFFFFF")
    set_cell_borders(kc)
    set_cell_borders(vc)
    kp = kc.paragraphs[0]
    kr = kp.add_run(k)
    kr.bold = True
    kr.font.size = Pt(10.5)
    kr.font.color.rgb = PRIMARY
    kr.font.name = "Calibri"
    vp = vc.paragraphs[0]
    vr = vp.add_run(v)
    vr.font.size = Pt(10.5)
    vr.font.color.rgb = BODY
    vr.font.name = "Calibri"

doc.add_page_break()

# ── ABSTRACT ─────────────────────────────────────────────────────────────
add_heading(doc, "Abstract", level=1)
add_para(doc,
    "The Integrated Patient Care Management System (IPCMS) is a web-based "
    "healthcare platform built with Python and Streamlit that consolidates "
    "the day-to-day clinical workflow into a single role-aware application. "
    "It serves three primary stakeholders — patients, doctors, and hospital "
    "administrators — and binds them through a unified MySQL data layer and "
    "two cooperating AI assistants."
)
add_para(doc,
    "The first AI assistant, SmartCare AI (ai_care.py), is a database-aware "
    "LangChain + Groq chatbot that can answer natural-language questions, "
    "summarise patient history, list available doctors, surface system "
    "statistics, and trigger an appointment-booking flow. The second is a "
    "voice-enabled chatbot (voice_chatbot.py) built on the Gemini Live API "
    "with Web Speech speech-to-text and text-to-speech, available as a "
    "floating widget on every page."
)
add_para(doc,
    "Beyond AI, the system implements end-to-end features including role-aware "
    "authentication with Google OAuth, self-service appointment booking with "
    "live slot conflict detection, an in-app pharmacy with admin-side CRUD on "
    "medicines, OCR-powered medical document analysis, automated PDF medical-"
    "record generation, and email-based doctor onboarding. The result is a "
    "working, demo-ready product that demonstrates how large language models "
    "can be safely embedded into a clinical workflow when they are scoped to "
    "the same data the user is currently looking at."
)
add_horizontal_rule(doc)

# ── 1. INTRODUCTION ─────────────────────────────────────────────────────
add_heading(doc, "1. Introduction", level=1)
add_heading(doc, "1.1 Background", level=2)
add_para(doc,
    "Healthcare delivery has long suffered from fragmented data, manual "
    "intake processes, and inconsistent channels of communication between "
    "patients, doctors, and hospital administrators. Even in urban centres, "
    "patient history is often split across multiple clinics, prescriptions "
    "live on paper, and onboarding a new doctor involves manual phone calls "
    "and shared passwords. The proliferation of large language models offers "
    "a new opportunity: AI can act as a connective layer that summarises, "
    "retrieves, and explains clinical information in natural language — but "
    "only if it is grounded in the same authoritative data the rest of the "
    "system uses."
)
add_heading(doc, "1.2 Problem Statement", level=2)
add_para(doc, "The project addresses four concrete problems:")
add_bullet(doc, "Fragmented patient records — notes live in different places "
                "for every doctor, making history hard to retrieve and harder "
                "to share.")
add_bullet(doc, "Manual appointment routing — reception staff read forms and "
                "assign slots by hand; slow, error-prone, and with no self-"
                "service option for the patient.")
add_bullet(doc, "No intelligent first response — patients Google symptoms "
                "or wait days for advice; doctors cannot scale 1-to-1 chat.")
add_bullet(doc, "Paper-based pharmacy and onboarding — medicine orders, "
                "refills, and doctor onboarding all run on paper and manual "
                "calls.")

add_heading(doc, "1.3 Objectives", level=2)
add_bullet(doc, "Provide a single web application with role-aware dashboards "
                "for patient, doctor, and admin.")
add_bullet(doc, "Wire two production-grade AI assistants — text and voice — "
                "into the same MySQL database that drives the rest of the app.")
add_bullet(doc, "Implement secure authentication, slot-aware booking, e-"
                "prescriptions, and email-based doctor onboarding.")
add_bullet(doc, "Provide an OCR pathway so patients and doctors can upload "
                "lab reports and have them read and summarised by AI.")

add_heading(doc, "1.4 Scope", level=2)
add_para(doc,
    "The system is built as a single Streamlit multi-page application backed "
    "by MySQL (with a SQLite fallback for laptops without a MySQL server). "
    "It supports three roles, four AI features, twelve end-user features, "
    "and a complete admin analytics layer. Production deployment is out of "
    "scope for this build; the focus is on demonstrating the architecture, "
    "the AI integration, and the complete feature loop."
)
add_horizontal_rule(doc)

# ── 2. SYSTEM ARCHITECTURE ───────────────────────────────────────────────
add_heading(doc, "2. System Architecture", level=1)
add_heading(doc, "2.1 High-Level Architecture", level=2)
add_para(doc,
    "The application follows a three-tier architecture: a Streamlit-rendered "
    "frontend, a Python backend that hosts both the business logic and the "
    "AI orchestrators, and a MySQL persistence layer. Two AI side-paths — "
    "the LangChain-based text assistant and the Gemini Live voice widget — "
    "share the same database connection pool used by the regular CRUD "
    "workflows."
)
add_callout(doc,
    "Architectural principle — grounded AI",
    "Every LLM call in this system is given the same DB slice the user is "
    "currently looking at. This is what makes 'book Dr. Sharma at 4 pm' "
    "actually work: the AI is not inventing availability, it is reading it."
)

add_heading(doc, "2.2 Component Diagram (logical)", level=2)
add_grid_table(doc,
    ["Tier", "Component", "Responsibility"],
    [
        ("Presentation", "Streamlit pages (app.py + pages/)",
         "Role-aware dashboards, forms, session state, custom CSS theme."),
        ("Presentation", "Voice widget (streamlit.components.v1.html)",
         "Always-on mic button pinned bottom-right; STT → LLM → TTS loop."),
        ("Application", "ai_care.py", "SmartCare AI text assistant, intent "
         "router, booking flow, OCR / medicine-image analysis."),
        ("Application", "auth.py", "Authentication, role detection, Google "
         "OAuth, password hashing (SHA-256)."),
        ("Application", "email_service.py", "SMTP HTML email for doctor "
         "credential delivery via Gmail."),
        ("Application", "db.py + db_config.py", "Schema builder, CRUD "
         "helpers, dialect-aware MySQL ↔ SQLite switching."),
        ("AI / LLM", "LangChain + ChatGroq (Groq SDK fallback)",
         "openai/gpt-oss-120b primary; llama-3.1-8b-instant fallback for "
         "short / latency-sensitive queries."),
        ("AI / LLM", "Gemini Live API + Groq vision models",
         "Voice loop, vision OCR for PDFs/images, medicine-image analysis."),
        ("Data", "MySQL 8 (PyMySQL) + SQLite (stdlib fallback)",
         "users, doctors, patients, slots, appointments, prescriptions, "
         "medicines, health_records, specialties."),
        ("External", "Google OAuth 2.0, Gmail SMTP, Groq API, Gemini API",
         "Identity, transactional mail, LLM inference, voice."),
    ],
    col_widths=[1.1, 2.4, 4.1]
)

add_heading(doc, "2.3 Module Layout", level=2)
add_para(doc, "The project is organised into the following modules:")
add_kv_table(doc, [
    ("app.py", "Streamlit entry point — public landing, login, signup, "
     "public chatbot, OAuth callback handling."),
    ("auth.py", "Hashing, seed users, role detection, Google OAuth "
     "handlers."),
    ("db.py", "Schema builder (dialect-aware), CRUD helpers, appointment "
     "booking with slot-conflict checks."),
    ("db_config.py", "Connection configuration (MySQL credentials and "
     "SQLite fallback)."),
    ("ai_care.py", "SmartCare AI — LangChain + Groq text assistant, "
     "appointment-booking flow, OCR, medicine-image analysis."),
    ("chatbot.py", "Public (pre-login) LLM chat panel for visitors."),
    ("voice_chatbot.py", "Floating voice widget using the Gemini Live API "
     "and Web Speech."),
    ("email_service.py", "SMTP HTML email for sending doctor credentials."),
    ("add_medicines_admin.py", "Admin-side medicine CRUD helper."),
    ("update_admin.py", "Admin-side user management helper."),
    ("pages/", "patient_dashboard.py, doctor_dashboard.py, "
     "admin_dashboard.py, shared_styles.py — the multi-page UI layer."),
    ("patient_care/", "Patient-care-specific helpers (vitals, prescriptions, "
     "PDF export)."),
    ("Medicine_Images/", "Static medicine catalogue assets."),
])

add_horizontal_rule(doc)

# ── 3. FEATURES ──────────────────────────────────────────────────────────
add_heading(doc, "3. Feature Set", level=1)
add_heading(doc, "3.1 Patient features", level=2)
add_bullet(doc, "Email + password sign-up and sign-in; Google OAuth entry "
                "point.")
add_bullet(doc, "Personal profile with medical history view.")
add_bullet(doc, "Browse doctors by specialty; book / reschedule / cancel "
                "appointments with live slot-conflict detection.")
add_bullet(doc, "AI symptom checker (text + vision) — describe a problem or "
                "upload a medicine photo and receive a structured analysis.")
add_bullet(doc, "Download prescription and visit summary as PDF.")
add_bullet(doc, "Floating voice chatbot available 24×7 on every page.")

add_heading(doc, "3.2 Doctor features", level=2)
add_bullet(doc, "Today's schedule and patient queue.")
add_bullet(doc, "AI-generated 5-bullet pre-visit brief per patient (history, "
                "recent vitals, current meds, red flags, suggested questions).")
add_bullet(doc, "Diagnosis entry, vitals, free-text notes.")
add_bullet(doc, "Write and edit prescriptions in-session.")
add_bullet(doc, "Manage availability slots.")
add_bullet(doc, "Password reset delivered by email on admin request.")

add_heading(doc, "3.3 Admin features", level=2)
add_bullet(doc, "Onboard new doctors with one click — credentials are auto-"
                "emailed.")
add_bullet(doc, "Activate / deactivate any user account.")
add_bullet(doc, "Specialty catalogue oversight; slot templates.")
add_bullet(doc, "Analytics dashboard: visits, revenue, doctor load — backed "
                "by pandas DataFrames rendered from the same MySQL DB.")
add_bullet(doc, "Reset any doctor's password.")
add_bullet(doc, "Full DB-backed audit view across the system.")

add_heading(doc, "3.4 Cross-cutting features", level=2)
add_bullet(doc, "Role-aware navigation — what a patient sees is never what a "
                "doctor sees.")
add_bullet(doc, "Single MySQL schema read by all three dashboards.")
add_bullet(doc, "Session state per role; chat history persisted in Streamlit "
                "session per user.")
add_bullet(doc, "Custom CSS theme (Inter / Plus Jakarta Sans, glassmorphism, "
                "medical slate palette).")
add_bullet(doc, "Responsive layout that collapses gracefully on narrow "
                "screens.")

add_horizontal_rule(doc)

# ── 4. AI IMPLEMENTATIONS (deep dive) ────────────────────────────────────
add_heading(doc, "4. AI Implementations", level=1)
add_para(doc,
    "Four distinct AI features are wired into the production application. "
    "Each is implemented as a self-contained module with a clear contract — "
    "the LLM receives a role-aware context payload, returns structured or "
    "free-text output, and the application layer is responsible for "
    "validating any action that mutates persistent state."
)

# 4.1 SmartCare AI
add_heading(doc, "4.1 SmartCare AI — in-app copilot (ai_care.py)", level=2)
add_para(doc,
    "SmartCare AI is the primary AI feature and the most heavily exercised. "
    "It is built on LangChain's ChatGroq wrapper around Groq's hosted LLMs. "
    "The default model is openai/gpt-oss-120b; if that model is unavailable "
    "or returns an error, the system falls back to llama-3.1-8b-instant, "
    "which trades depth for latency and is preferable for short, structured "
    "queries."
)
add_heading(doc, "Architecture", level=3)
add_bullet(doc, "Lazy-loaded ChatGroq client keyed on GROQ_API_KEY from the "
                ".env file. Missing keys produce a friendly fallback message "
                "instead of a hard crash.")
add_bullet(doc, "Intent router: keyword-matching pre-LLM layer handles "
                "book appointment, list doctors, my appointments, my vitals, "
                "admin stats, and medicine lookup without an LLM call. This "
                "keeps deterministic lookups fast and free.")
add_bullet(doc, "Live DB context: before any LLM call, _build_db_context() "
                "summarises the relevant rows from MySQL into a short text "
                "payload — counts, top doctors, the calling user's "
                "appointments and vitals. The LLM is grounded in this "
                "context rather than allowed to invent.")
add_bullet(doc, "Booking flow trigger: if the LLM decides the user wants to "
                "book, it includes the sentinel token [OPEN_BOOKING_FORM] "
                "in its response. The application layer strips the token "
                "and surfaces a four-step Streamlit booking form (doctor → "
                "date → slot → confirm) with db.check_slot_conflict() "
                "guarding every step.")
add_bullet(doc, "Booking write-back: confirmed bookings go through "
                "db.book_appointment() — never directly through the LLM — "
                "so the data layer remains the single source of truth.")
add_heading(doc, "System prompt (abridged)", level=3)
add_para(doc,
    "“You are SmartCare AI, an advanced medical AI assistant integrated into "
    "PCMHS. Current User Info: Name=…, Role=…. LIVE DATABASE GROUND TRUTH "
    "CONTEXT: <DB snapshot>. 1. Assist with general medical knowledge and "
    "platform guidance. 2. Use the live DB context for accurate answers "
    "about doctors, appointments, vitals, and system stats. 3. ALWAYS state "
    "that you are an AI assistant, NOT a doctor. 4. For emergencies, direct "
    "users to local emergency services (108 / 911). 5. Keep responses "
    "structured, concise, empathetic. 6. Always close with an offer of "
    "further help. 7. If the user wants to book, include [OPEN_BOOKING_FORM].”",
    italic=True, color=SLATE
)

# 4.2 Voice chatbot
add_heading(doc, "4.2 Floating voice chatbot (voice_chatbot.py)", level=2)
add_para(doc,
    "The voice chatbot is an always-on widget rendered via "
    "streamlit.components.v1.html() and pinned bottom-right on every "
    "page. It uses the browser's Web Speech API for speech-to-text and "
    "text-to-speech, and the Google Gemini Live API for the model. If the "
    "mic fails or HTTPS is unavailable, it degrades to a text-only chat "
    "panel without losing the AI backend."
)
add_bullet(doc, "STT — Web Speech API listens on a user gesture; the captured "
                "transcript is posted to the model.")
add_bullet(doc, "LLM — Gemini Live API in the browser; Groq is wired as a "
                "secondary fallback for environments where Gemini is "
                "blocked.")
add_bullet(doc, "TTS — Web Speech API speaks the model response aloud, with "
                "a mute toggle for users who prefer silent reading.")
add_bullet(doc, "Privacy — audio capture is never persisted; only the "
                "transcript round-trips through the model.")

# 4.3 Public chatbot
add_heading(doc, "4.3 Public (pre-login) LLM chat (chatbot.py)", level=2)
add_para(doc,
    "On the landing page, before authentication, a small ‘ask anything’ panel "
    "is exposed to visitors. It uses ChatGroq with LangChain Core for "
    "session state. Because the visitor is unauthenticated, the model is "
    "explicitly forbidden from making up doctor availability or system "
    "statistics; it answers FAQs about the product and the AI features "
    "instead."
)

# 4.4 OCR + vision
add_heading(doc, "4.4 OCR document analysis & medicine-image analysis", level=2)
add_para(doc,
    "Two vision-driven features live inside ai_care.py as "
    "analyze_document_ocr() and analyze_medicine_image(). They share the "
    "same backend pattern: base64-encode the upload, send it to a Groq "
    "vision-capable model (qwen/qwen3.6-27b by default, with dynamic "
    "discovery of any available multimodal model), and return a structured "
    "analysis."
)
add_heading(doc, "analyze_document_ocr()", level=3)
add_bullet(doc, "Images (jpg/png/webp/bmp) — sent directly to the vision "
                "model for OCR text extraction + analysis.")
add_bullet(doc, "PDFs — first pass through pdfplumber to extract the text "
                "layer; embedded images are cropped and individually OCR'd; "
                "if no text layer exists, the first page is rendered to a "
                "PNG at 200 DPI and OCR'd as a full-page image.")
add_bullet(doc, "A second LLM call (llama-3.1-8b-instant) then summarises "
                "the extracted text under five headings: Document Summary, "
                "Medical Information Detected, Important Findings, Key "
                "Data Points, Recommendations.")
add_heading(doc, "analyze_medicine_image()", level=3)
add_bullet(doc, "Patient uploads a photo of a medicine strip / bottle; the "
                "vision model returns: Medicine Name & Brand, Description, "
                "Medical Uses, Benefits, Limitations, Dosage, Side Effects, "
                "Contraindications, Drug Interactions, Storage, and a "
                "Medical Disclaimer.")
add_bullet(doc, "Designed for low-literacy users who can't read the printed "
                "label — the model is explicitly told to ask for a clearer "
                "photo if the medicine cannot be identified.")
add_para(doc,
    "Both vision routines strip <think>...</think> tags emitted by "
    "reasoning models before returning the final text, and never mutate the "
    "database — they are read-only analyses of the user's own upload."
)

# 4.5 Doctor pre-visit brief
add_heading(doc, "4.5 Doctor pre-visit brief", level=2)
add_para(doc,
    "Before each consultation, the doctor dashboard renders a 5-bullet "
    "patient brief generated by the same _build_db_context() used by the "
    "chatbot. The brief surfaces key history, recent vitals, current "
    "medications, red flags, and suggested questions to ask. Because the "
    "context is computed live from the same DB the application reads from, "
    "doctors can trust the brief the moment they see it."
)

# 4.6 AI safety
add_heading(doc, "4.6 AI safety and guardrails", level=2)
add_bullet(doc, "Every response opens with — or contains — a reminder that "
                "the assistant is an AI, not a doctor.")
add_bullet(doc, "Medical-emergency intent is intercepted in the system prompt "
                "and redirected to local emergency numbers (108 / 911).")
add_bullet(doc, "All persistent writes — appointments, prescriptions, "
                "password resets — go through typed db.* helpers, never "
                "through the LLM.")
add_bullet(doc, "The DB context payload sent to the model is restricted to "
                "the role of the calling user; a patient never sees another "
                "patient's vitals through the AI.")
add_bullet(doc, "Vision analyses are scoped to the user's own upload; "
                "nothing is written back to the database from image OCR.")

add_horizontal_rule(doc)

# ── 5. TECH STACK ────────────────────────────────────────────────────────
add_heading(doc, "5. Technology Stack", level=1)
add_grid_table(doc,
    ["Layer", "Technology", "Version", "Role in the system"],
    [
        ("Frontend / UI", "Streamlit",        "≥ 1.30",
         "Whole web app shell — pages, widgets, session state, file uploads, "
         "custom CSS theme."),
        ("Backend logic", "Python",           "3.9+",
         "Application logic, business rules, all AI orchestration."),
        ("Database",      "MySQL via PyMySQL", "≥ 1.1",
         "Primary persistence for users, doctors, slots, visits, "
         "prescriptions, medicines."),
        ("Database",      "SQLite (stdlib)",  "—",
         "Transparent fallback when MySQL is not available."),
        ("AI / LLM",      "LangChain + ChatGroq", "core ≥ 0.2, groq ≥ 0.1",
         "Primary AI brain — SmartCare AI text assistant."),
        ("AI / LLM",      "groq (SDK)",       "≥ 0.37",
         "Direct fallback path and vision-model client."),
        ("AI / LLM",      "Gemini Live API",  "—",
         "Voice widget model + browser-side STT/TTS."),
        ("Auth",          "Google OAuth 2.0", "—",
         "Optional OAuth alongside email/password."),
        ("Auth",          "hashlib (stdlib)", "—",
         "SHA-256 password hashing — zero-install demo path."),
        ("Email",         "smtplib (stdlib)", "—",
         "SMTP delivery of doctor credentials via Gmail app password."),
        ("Reports",       "ReportLab",        "≥ 4.0",
         "PDF medical-record generation."),
        ("Reports",       "pdfplumber",       "≥ 0.10",
         "Parsing uploaded lab reports before AI summarisation."),
        ("Reports",       "python-pptx",      "—",
         "Generates the project showcase deck (generate_ppt.py)."),
        ("Utilities",     "pandas",           "≥ 2.0",
         "Admin analytics — visit counts, revenue, doctor load."),
        ("Utilities",     "Pillow",           "≥ 10",
         "Logo and favicon rendering."),
        ("Utilities",     "requests",         "≥ 2.31",
         "Google OAuth HTTP flow."),
        ("Config",        "python-dotenv",    "≥ 1.0",
         "Loads .env into os.environ; secrets never live in source."),
        ("Crypto",        "cryptography",     "≥ 41",
         "Required transitively by PyMySQL for caching_sha2_password."),
    ],
    col_widths=[1.3, 1.8, 1.0, 3.5]
)

add_horizontal_rule(doc)

# ── 6. DATABASE DESIGN ───────────────────────────────────────────────────
add_heading(doc, "6. Database Design", level=1)
add_para(doc,
    "The schema is built once at startup by db.py and is dialect-aware: it "
    "emits AUTO_INCREMENT for MySQL and AUTOINCREMENT for SQLite, BOOLEAN for "
    "MySQL and INTEGER for SQLite. The application code talks to a single "
    "set of CRUD helpers regardless of which backend answered the connection."
)
add_grid_table(doc,
    ["Table", "Purpose", "Key columns"],
    [
        ("users",         "All accounts, role-aware", "id, email, password_hash, "
         "full_name, role (patient/doctor/admin), is_active, oauth_provider"),
        ("doctors",       "Doctor profile extension", "id, user_id, specialty, "
         "consultation_fee, experience_years, bio"),
        ("patients",      "Patient profile extension", "id, user_id, dob, gender, "
         "blood_group, address"),
        ("specialties",   "Catalogue of medical specialties", "id, name"),
        ("slots",         "Doctor availability windows", "id, doctor_id, weekday, "
         "start_time, end_time"),
        ("appointments",  "Booked visits", "id, patient_id, doctor_id, "
         "scheduled_date, start_time, end_time, status, reason"),
        ("prescriptions", "Doctor-issued medications", "id, appointment_id, "
         "medicine_id, dosage, duration, notes"),
        ("health_records","Vitals captured during visits", "id, patient_id, "
         "heart_rate, blood_pressure, pulse_oximetry, ejection_fraction, "
         "cardiac_output, diagnosis, notes"),
        ("medicines",     "Pharmacy catalogue", "id, name, category, unit_price, "
         "stock_quantity, image_path"),
    ],
    col_widths=[1.4, 2.0, 4.2]
)

add_horizontal_rule(doc)

# ── 7. UI / UX ───────────────────────────────────────────────────────────
add_heading(doc, "7. User Interface & Experience", level=1)
add_heading(doc, "7.1 Design system", level=2)
add_bullet(doc, "Theme — medical white background with glassmorphism overlays "
                "and a soft slate / sky / teal palette.")
add_bullet(doc, "Typography — Inter and Plus Jakarta Sans via Google Fonts.")
add_bullet(doc, "Backgrounds — base64-encoded image + soft overlay per page; "
                "consistent across roles.")
add_bullet(doc, "Components — cards, chips, pill buttons, sticky header, and "
                "the floating voice button.")
add_bullet(doc, "Responsiveness — wide layout by default; cards collapse "
                "gracefully on narrow screens.")

add_heading(doc, "7.2 Per-role page structure", level=2)
add_grid_table(doc,
    ["Page", "File", "Highlights"],
    [
        ("Landing / login", "app.py",
         "Public hero, sign-up form, Google OAuth button, public chatbot."),
        ("Patient dashboard", "pages/patient_dashboard.py",
         "Profile + history, doctor browse, AI Care, prescriptions, PDF "
         "downloads, voice widget."),
        ("Doctor dashboard", "pages/doctor_dashboard.py",
         "Today's queue, AI pre-visit brief, vitals entry, prescription "
         "writer, slot management."),
        ("Admin dashboard", "pages/admin_dashboard.py",
         "Onboard doctors, user management, pharmacy CRUD, analytics "
         "with pandas, password reset."),
        ("Shared styles", "pages/shared_styles.py",
         "Centralised CSS and theme tokens, medical banner component."),
    ],
    col_widths=[1.6, 2.4, 3.6]
)

add_horizontal_rule(doc)

# ── 8. SECURITY ──────────────────────────────────────────────────────────
add_heading(doc, "8. Security & Privacy", level=1)
add_bullet(doc, "Passwords are stored as SHA-256 hashes in the users table; "
                "no plaintext leaves the application.")
add_bullet(doc, "Role-based access control is enforced server-side; the "
                "client-side role is never trusted alone.")
add_bullet(doc, "Google OAuth redirect URIs are pinned to localhost:8501 to "
                "prevent open-redirect issues in development.")
add_bullet(doc, "Secrets (GROQ_API_KEY, SMTP password, OAuth client secret) "
                "are loaded from .env via python-dotenv and never committed.")
add_bullet(doc, "AI calls receive only the slice of the DB that the calling "
                "role is allowed to see — patients cannot retrieve another "
                "patient's vitals through the chatbot.")
add_bullet(doc, "Vision analyses are scoped to the user's own upload; the "
                "model is not asked to identify third parties.")
add_bullet(doc, "E-prescription PDFs are generated client-side via ReportLab "
                "and never sent to a third-party service.")

add_horizontal_rule(doc)

# ── 9. DIFFICULTIES ──────────────────────────────────────────────────────
add_heading(doc, "9. Difficulties Encountered", level=1)
add_grid_table(doc,
    ["Difficulty", "What went wrong", "Fix"],
    [
        ("MySQL not running locally",
         "Most laptops don't have MySQL installed; first launch crashed.",
         "Dynamic db_config with a SQLite fallback; UI prompt to enter MySQL "
         "creds when the user wants to switch."),
        ("Streamlit reruns everything",
         "Any click reruns the whole script, causing duplicate LLM calls.",
         "st.session_state keys per role and per user hold chat history and "
         "form state."),
        ("OAuth redirect in dev",
         "Google OAuth requires exact redirect URIs; 127.0.0.1 vs localhost "
         "matters.",
         "Hard-coded the redirect URI in auth.py; documented the localhost "
         "requirement."),
        ("AI cost & latency",
         "Long LLM calls blocked the UI; cost grew with concurrent users.",
         "st.spinner during calls; smaller prompts; llama-3.1-8b-instant "
         "fallback for short queries."),
        ("Voice in the browser",
         "Web Speech requires a user gesture and HTTPS (or localhost).",
         "Graceful text-only fallback when mic or speech APIs are "
         "unavailable."),
        ("Schema drift MySQL ↔ SQLite",
         "AUTO_INCREMENT vs AUTOINCREMENT, BOOL vs BOOLEAN, VARCHAR lengths.",
         "Dialect-aware schema builder in db.py centralises the differences."),
        ("Vision model availability",
         "Vision-capable model IDs change frequently on Groq.",
         "Dynamic discovery via client.models.list() with a curated default "
         "list as fallback."),
    ],
    col_widths=[1.7, 2.7, 3.2]
)

add_horizontal_rule(doc)

# ── 10. TESTING ──────────────────────────────────────────────────────────
add_heading(doc, "10. Testing & Verification", level=1)
add_para(doc,
    "Testing is performed manually against the running Streamlit app. The "
    "following checklist is exercised end-to-end before any release."
)
add_bullet(doc, "Authentication — sign-up, sign-in, Google OAuth round-trip, "
                "invalid credentials, role detection.")
add_bullet(doc, "Booking — happy path, slot conflict, double-booking race, "
                "cancellation, doctor availability edit.")
add_bullet(doc, "AI assistant — intent router paths (book, list, vitals, "
                "admin stats), LLM fallback when GROQ_API_KEY is missing.")
add_bullet(doc, "OCR — image upload, scanned PDF with no text layer, multi-"
                "page PDF, embedded-image extraction.")
add_bullet(doc, "Voice widget — mic capture, transcript round-trip, mute "
                "toggle, text-only fallback.")
add_bullet(doc, "Doctor onboarding — admin creates doctor → email arrives → "
                "doctor signs in with the temp password → resets it.")
add_bullet(doc, "Pharmacy — admin CRUD, patient browse, low-stock visibility "
                "in the AI context.")
add_bullet(doc, "Admin analytics — visit counts, revenue, doctor load render "
                "correctly with both MySQL and SQLite backends.")

add_horizontal_rule(doc)

# ── 11. FUTURE WORK ──────────────────────────────────────────────────────
add_heading(doc, "11. Future Work", level=1)
add_bullet(doc, "Multi-language chatbot — Hindi, Spanish, and regional "
                "languages.")
add_bullet(doc, "Doctor-side calendar sync (Google Calendar, iCal).")
add_bullet(doc, "Insurance and billing module with claim submission.")
add_bullet(doc, "Mobile-first layout — current build is desktop-tuned.")
add_bullet(doc, "LLM cost telemetry and rate-limiting per role.")
add_bullet(doc, "MySQL connection pooling for concurrent doctors.")
add_bullet(doc, "Stronger PII safety: no raw history in prompts, redaction "
                "layer between DB context and LLM input.")

add_horizontal_rule(doc)

# ── 12. CONCLUSION ───────────────────────────────────────────────────────
add_heading(doc, "12. Conclusion", level=1)
add_para(doc,
    "IPCMS demonstrates that an AI-assisted clinical workflow does not need "
    "to be magical to be useful. By grounding every LLM call in the same "
    "MySQL state that the rest of the application reads from, the system "
    "stays auditable, deterministic where it matters, and helpful where it "
    "can be — appointment booking that actually checks availability, "
    "doctor briefs that actually reflect the patient's last visit, OCR that "
    "actually reads the uploaded PDF. The result is a working product that "
    "moves the right information to the right person, faster, without "
    "replacing the human in the loop."
)
add_para(doc,
    "The codebase is intentionally written to be read: each module is small, "
    "the AI orchestration is centralised in ai_care.py, the data layer is "
    "centralised in db.py, and the UI is composed from reusable shared "
    "styles. This makes the system approachable as a teaching artefact and "
    "as a starting point for a more ambitious clinical platform."
)

# ── REFERENCES ───────────────────────────────────────────────────────────
add_heading(doc, "References", level=1)
add_bullet(doc, "Streamlit documentation — https://docs.streamlit.io")
add_bullet(doc, "LangChain documentation — https://python.langchain.com")
add_bullet(doc, "Groq Cloud — https://console.groq.com")
add_bullet(doc, "Google Gemini API — https://ai.google.dev")
add_bullet(doc, "PyMySQL — https://pymysql.readthedocs.io")
add_bullet(doc, "ReportLab — https://www.reportlab.com")
add_bullet(doc, "pdfplumber — https://github.com/jsvine/pdfplumber")
add_bullet(doc, "python-pptx — https://python-pptx.readthedocs.io")
add_bullet(doc, "Web Speech API — MDN Web Docs")

# ── Save ─────────────────────────────────────────────────────────────────
out = r"E:\Integrated_Patient_Care_Management_System\IPCMS_Project_Report.docx"
doc.save(out)
print(f"Saved {out}")
